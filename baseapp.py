import csv
import json
import pickle
import random
import streamlit as st
import pandas as pd
import flask
from flask import session
from streamlit_extras.grid import grid
from st_aggrid import AgGrid, GridOptionsBuilder
import warnings
import openpyxl
import re
from docx import Document
from collections import defaultdict
from io import BytesIO


warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
st.set_page_config(layout="wide")


def to_excel_custom_wTechstring(df, techstring_out):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Data', index=False)
        # Write the techstring as a DataFrame in the second sheet
        pd.DataFrame({'TechString': [techstring_out]}).to_excel(writer, sheet_name='Summary', index=False)
    return output.getvalue()


def mainpage():

    maintab, manualvartab, debugtab = st.tabs(['Main', 'ManualVar', 'Debug'])

    with maintab:

        col_starty, col_startyq, dash, col_endy, col_endq = st.columns([3, 1, 4, 3, 1])
        with col_starty:
            st.number_input('Start', 2000, 2024, key='starty', value=2010)
        with col_startyq:
            st.number_input('Q', 1, 4, key='startq', value=1)
        with dash:
            st.subheader('   ')
        with col_endy:
            st.number_input('Last Known', 2000, 2024, key='endy', value=2020)
        with col_endq:
            st.number_input('Q', 1, 4, key='endq', value=4)

        my_grid = grid([0.5, 1, 4, 6, 6, 2, 1.33], vertical_align="bottom")

        if st.session_state.get('sidebar_done', False):
            variables = st.session_state['formulas']
            print(variables)
            dbkeys = st.session_state['dbkeys']
            excel_defaults = st.session_state['excel_defaults_onecountry']
            most_frequent_country_code = excel_defaults['reporter'].mode().iloc[
                0]  # set this as the default option for ccode dropdowns

            try:
                country_dictionary = pd.read_csv('streamlit_assets/country_dictionary.csv')
                most_frequent_country_name = \
                    country_dictionary[country_dictionary['code'] == most_frequent_country_code]['name'].iloc[0]
            except Exception as e:
                print(e)

            for col in excel_defaults.columns:
                if col != 'eq':  # Skip the 'eq' column
                    if pd.api.types.is_string_dtype(excel_defaults[col]):  # Check if the column contains string values
                        excel_defaults[col] = excel_defaults[col].str.lower()

            dict_indicators_units = dbkeys.groupby('reporter').apply(
                lambda df: {
                    indicator: list(df[df['indicator'] == indicator]['unit'].unique())
                    for indicator in df['indicator'].unique()
                }, include_groups=False
            ).to_dict()

            dict_units_lids = dbkeys.groupby(by=['reporter', 'indicator', 'unit'])['lid'].apply(list).to_dict()

            dict_final = {}
            for reporter, indicators_units in dict_indicators_units.items():
                dict_final[reporter] = {}
                for indicator, units in indicators_units.items():
                    dict_final[reporter][indicator] = {}
                    for unit in units:
                        lids = dict_units_lids.get((reporter, indicator, unit), [])
                        dict_final[reporter][indicator][unit] = lids

            for value in variables['x']:
                var_blocked = my_grid.checkbox('block', key=f'block_{value}', label_visibility='collapsed')
                my_grid.text(value)

                if not var_blocked:
                    try:

                        default_indicator = None
                        default_unit = None
                        default_lid = None

                        try:
                            # default_reporter = most_frequent_country_name
                            default_reporter = None
                        except Exception as e:
                            print(e)
                            default_reporter = None

                        if not excel_defaults.empty:
                            default_corresp_values = excel_defaults[excel_defaults['variable'] == value]
                            if not default_corresp_values.empty:
                                default_row = default_corresp_values.iloc[0]
                                default_lid = default_row['lid']

                                # Find corresponding reporter, indicator, and unit
                                for reporter, indicators_units in dict_indicators_units.items():
                                    for indicator, units in indicators_units.items():
                                        for unit in units:
                                            if default_lid in dict_units_lids.get((reporter, indicator, unit), []):
                                                default_reporter = reporter
                                                default_indicator = indicator
                                                default_unit = unit
                                                break
                                        if default_reporter:
                                            # break
                                            pass  # check later if this is app breaking
                                    if default_reporter:
                                        # break
                                        pass  # check later if this is app breaking

                        if default_lid:
                            my_grid.text(f"{default_reporter}")
                            st.session_state[f"country_{value}"] = default_reporter

                            my_grid.text(f"{default_indicator}")
                            st.session_state[f"indicator_{value}"] = default_indicator

                            my_grid.text(f"{default_unit}")
                            st.session_state[f"unit_{value}"] = default_unit

                            my_grid.text(f"{default_lid}")
                            st.session_state[f"lid_{value}"] = default_lid
                        else:
                            try:
                                index_country_mapping = dbkeys[dbkeys['reporter'] == most_frequent_country_name].index[0]
                            except IndexError:
                                print(f'Default country position not found. Reverting to default.')
                                index_country_mapping = 0  # Fallback index if the default country is not found

                            selected_country = my_grid.selectbox('Select Country', options=dbkeys['reporter'].unique(),
                                                                 key=f'country_{value}', index=int(index_country_mapping))
                            if selected_country:
                                available_indicators = dbkeys[dbkeys['reporter'] == selected_country]['indicator'].unique()
                                selected_indicator = my_grid.selectbox('Select Indicator', options=available_indicators,
                                                                       key=f'indicator_{value}')

                                if selected_indicator:
                                    available_units = \
                                        dbkeys[(dbkeys['reporter'] == selected_country) & (
                                                dbkeys['indicator'] == selected_indicator)][
                                            'unit'].unique()
                                    selected_unit = my_grid.selectbox('Select Unit', options=available_units,
                                                                      key=f'unit_{value}', index=None)

                                    if selected_unit:
                                        lids = dict_final.get(selected_country, {}).get(selected_indicator, {}).get(
                                            selected_unit,
                                            [])

                                        if lids:
                                            my_grid.text(*lids)
                                            st.session_state[f"lid_{value}"] = lids[0]
                                        else:
                                            my_grid.text("NA")
                                            st.session_state[f"lid_{value}"] = "NA"
                                    else:
                                        my_grid.text("Unit?")
                                        # st.session_state[f"unit_{value}"] = "Unit?"
                                else:
                                    my_grid.text("Please select an indicator.")
                                    # st.session_state[f"indicator_{value}"] = "Please select an indicator."
                            else:
                                my_grid.text("Please select a country.")
                                # st.session_state[f"country_{value}"] = "Please select a country."
                    except ValueError as e:
                        print(e)
                else:
                    my_grid.text(f"USERVAR")
                    st.session_state[f"country_{value}"] = "USERVAR"

                    my_grid.text(f"USERVAR")
                    st.session_state[f"indicator_{value}"] = "USERVAR"

                    my_grid.text(f"USERVAR")
                    st.session_state[f"unit_{value}"] = "USERVAR"

                    my_grid.text(f"USERVAR")
                    st.session_state[f"lid_{value}"] = "USERVAR"

                if not var_blocked:
                    if value in ['ir', 'eu_ir', 'loans_hh', 'pub_debt']:
                        my_grid.selectbox('r', options=['avg', 'eop'], index=1,
                                          key=f'periodflag_{value}', label_visibility='hidden')
                    else:
                        my_grid.selectbox('eop', options=['avg', 'eop'], index=0,
                                          key=f'periodflag_{value}', label_visibility='hidden')

                else:
                    my_grid.text(f"USERVAR")
                    st.session_state[f"periodflag_{value}"] = "USERVAR"

            st.header('')
            download_grid = grid([1, 3], [5, 1])

            generate_alldata_button = download_grid.button('Generate')
            download_randomizer_tickbox = download_grid.checkbox('Replace undeclared lids with random database entries? (excludes blocked and default)')

            if generate_alldata_button:
                alldata = []
                if not download_randomizer_tickbox:
                    for value in variables['x']:
                        country = st.session_state.get(f"country_{value}", "NA")
                        indicator = st.session_state.get(f"indicator_{value}", "NA")
                        unit = st.session_state.get(f"unit_{value}", "NA")
                        lid = st.session_state.get(f"lid_{value}", "NA")
                        periodflag = st.session_state.get(f"periodflag_{value}", "NA")

                        alldata.append({
                            'Variable': value,
                            'Country': country,
                            'Indicator': indicator,
                            'Unit': unit,
                            'LID': lid,
                            'PeriodFlag': periodflag
                        })

                else:
                    # Randomizing within country lids only!
                    forrand_country_chosen = most_frequent_country_name
                    forrand_country_db = dbkeys[dbkeys['reporter'] == forrand_country_chosen]
                    forrand_country_lids = forrand_country_db['lid'].unique()

                    for value in variables['x']:
                        country = st.session_state.get(f"country_{value}", forrand_country_chosen)
                        indicator = st.session_state.get(f"indicator_{value}", "random")
                        unit = st.session_state.get(f"unit_{value}", "random")
                        lid = st.session_state.get(f"lid_{value}", random.choice(forrand_country_lids))
                        periodflag = st.session_state.get(f"periodflag_{value}", "NA")

                        alldata.append({
                            'Variable': value,
                            'Country': country,
                            'Indicator': indicator,
                            'Unit': unit,
                            'LID': lid,
                            'PeriodFlag': periodflag
                        })

                # Create a DataFrame from the collected data
                df_alldata = pd.DataFrame(alldata)
                st.dataframe(df_alldata, use_container_width=True)

                starty_out = int(st.session_state['starty'])
                startq_out = int(st.session_state['startq'])
                endy_out = int(st.session_state['endy'])
                endq_out = int(st.session_state['endq'])
                techstring_out = f"{starty_out}Q{startq_out}//{endy_out}Q{endq_out}"

                # Print or log the result
                print(techstring_out)

                xlsx_data = to_excel_custom_wTechstring(df_alldata, techstring_out)
                st.download_button(label="Download Data", data=xlsx_data, file_name="streamlit_out.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        with manualvartab:

            @st.cache_data
            def generate_manual_table_data(input_num_columns, start_year, start_quarter, end_year, end_quarter):
                empty_columns = [f'Variable_name_{i + 1}' for i in range(input_num_columns)]

                dates = []
                for year in range(start_year, end_year + 1):
                    if year == start_year:
                        quarter_start = start_quarter
                    else:
                        quarter_start = 1
                    if year == end_year:
                        quarter_end = end_quarter
                    else:
                        quarter_end = 4
                    for quarter in range(quarter_start, quarter_end + 1):
                        dates.append(f"{year}Q{quarter}")

                data = pd.DataFrame({'date': dates})
                for col in empty_columns:
                    data[col] = ""

                return data

            def display_editable_grid(data):
                gb = GridOptionsBuilder.from_dataframe(data)
                gb.configure_default_column(editable=True)
                grid_options = gb.build()
                return AgGrid(data, gridOptions=grid_options, editable=True, fit_columns_on_grid_load=True)

            with st.expander("Manual Variable Table"):
                st.subheader('Declare your non-DB variables here')
                st.text('You can copypaste the data in from Excel')
                num_columns = st.number_input('How many additional columns to add?', min_value=1, max_value=10, step=1)

                if 'table_data' not in st.session_state:
                    st.session_state.table_data = None
                if 'freeze_updates' not in st.session_state:
                    st.session_state.freeze_updates = False

                manualvar_generate_button = st.button('Generate blank table')

                if manualvar_generate_button or (
                        st.session_state.table_data is not None and st.session_state.freeze_updates):
                    if manualvar_generate_button:
                        st.session_state.freeze_updates = True
                        st.session_state.manualvar_generate_button = False
                        st.session_state.table_data = generate_manual_table_data(
                            num_columns,
                            st.session_state['starty'],
                            st.session_state['startq'],
                            st.session_state['endy'],
                            st.session_state['endq']
                        )

                    grid_response = display_editable_grid(st.session_state.table_data)

                    manualvarfinal_generate_button = st.button('Generate final table')

                    if manualvarfinal_generate_button:
                        st.session_state.freeze_updates = False
                        updated_data = grid_response['data']
                        st.session_state.table_data = updated_data
                        st.download_button(label="Download CSV", data=updated_data.to_csv(index=False),
                                           file_name="generated_data.csv")
                        st.error('IMPORTANT: do not forget to RENAME THE VARIABLES to their proper names before using this dataframe!')

        with debugtab:
            randomize_button = st.button('feeling lucky')
            # st.session_state
            if randomize_button:
                uservar_lid_pairs = []
                for key, value in st.session_state.items():
                    if key.startswith('lid_'):
                        uservar = key.replace('lid_', '')
                        lid = value
                        # uservar = st.session_state.get(f'x_{key[0:]}', 'NA')  # Extract corresponding uservar
                        uservar_lid_pairs.append((uservar, lid))

                if uservar_lid_pairs:
                    with open('streamlit_assets/uservar_lid_pairs.csv', 'w', newline='') as csvfile:
                        csvwriter = csv.writer(csvfile)
                        csvwriter.writerow(['value', 'lid'])  # Write header
                        csvwriter.writerows(uservar_lid_pairs)  # Write data

                    st.dataframe(uservar_lid_pairs)

                    with open('streamlit_assets/uservar_lid_pairs.csv', 'rb') as file:
                        st.download_button('Download CSV', file, file_name='lid_uservar_pairs.csv')

                out_blocked = []
                for key, value in st.session_state.items():
                    if key.startswith('block_'):
                        uservar = key.replace('block_', '')
                        block = value
                        out_blocked.append((uservar, block))

                if out_blocked:
                    with open('out_blocked.csv', 'w', newline='') as csvfile:
                        csvwriter = csv.writer(csvfile)
                        csvwriter.writerow(['value', 'block'])  # Write header
                        csvwriter.writerows(out_blocked)  # Write data

                    st.success('CSV file with lid-uservar pairs generated!')
                    st.dataframe(out_blocked)

                    # Provide download link in streamlit
                    with open('out_blocked.csv', 'rb') as file:
                        st.download_button('Download CSV', file, file_name='out_blocked.csv')

                out_periodflag = []
                for key, value in st.session_state.items():
                    if key.startswith('periodflag_'):
                        uservar = key.replace('periodflag_', '')
                        periodflag = value
                        out_periodflag.append((uservar, periodflag))

                if out_periodflag:
                    with open('out_periodflag.csv', 'w', newline='') as csvfile:
                        csvwriter = csv.writer(csvfile)
                        csvwriter.writerow(['value', 'out_periodflag'])  # Write header
                        csvwriter.writerows(out_periodflag)  # Write data

                    st.dataframe(out_periodflag)

                    with open('out_periodflag.csv', 'rb') as file:
                        st.download_button('Download CSV', file, file_name='out_periodflag.csv')


def sidepage():

    def parse_formulas(formulas_file):
        def convert_lag_notation(formula):
            # Convert lag(x, number) to `lag(x, number)`
            pattern1 = r"lag\((\w+),\s*(-?\d+)\)"
            formula = re.sub(pattern1, r"`lag(\1,\2)`", formula)

            # Handle lags written as x(-number)
            pattern2 = r"(\w+)\(\s*-\s*(\d+)\s*\)"
            formula = re.sub(pattern2, r"`lag(\1,\2)`", formula)  # Convert to lag format without '-'

            return formula

        def convert_log_notation(formula):
            pattern = r"log\((\w+)\)"
            return re.sub(pattern, r"`log(\1)`", formula)

        def convert_exp_notation(formula):
            pattern = r"exp\((\w+)\)"
            return re.sub(pattern, r"`exp(\1)`", formula)

        # Check if file is passed correctly
        print("Formulas file received:", formulas_file)

        # Try opening the docx file
        try:
            doc = Document(formulas_file)
            formulas = [para.text.lower() for para in doc.paragraphs if para.text.strip()]
            print("Extracted formulas:", formulas)
        except Exception as e:
            print("Error reading file:", e)
            return None, None

        # Ensure formulas are processed
        if not formulas:
            print("No formulas found in document.")
            return None, None

        # Process formulas
        formulas = [formula.replace("=", "~").replace(" ", "") for formula in formulas]
        formulas = [convert_lag_notation(formula) for formula in formulas]
        formulas = [convert_log_notation(formula) for formula in formulas]
        formulas = [convert_exp_notation(formula) for formula in formulas]

        # Check processed formulas
        print("Processed formulas:", formulas)

        # Extract variables
        word_vector = [
            re.findall(r"`[^`]+`|\b(?:[a-zA-Z_]+\d*[a-zA-Z_]*|\d+[a-zA-Z_]+(?:\.\d+)?|[a-zA-Z_]+(?:\.\d+)?)\b", formula)
            for formula in formulas
        ]
        print("Extracted word vector:", word_vector)
        word_vector = list(set([item for sublist in word_vector for item in sublist]))
        print("Unique word vector:", word_vector)

        # Create DataFrame
        vars_n_DBvars_fetched = defaultdict(list)
        for var in word_vector:
            if var.startswith("`lag"):
                vars_n_DBvars_fetched['lag'].append(var)
                # Extract the variable inside lag and add to 'x'
                inner_var = re.search(r"`lag\(([^,]+),", var)
                if inner_var:
                    vars_n_DBvars_fetched['x'].append(inner_var.group(1))
            elif var.startswith("`log"):
                vars_n_DBvars_fetched['log'].append(var)
                # Extract the variable inside log and add to 'x'
                inner_var = re.search(r"`log\(([^)]+)\)", var)
                if inner_var:
                    vars_n_DBvars_fetched['x'].append(inner_var.group(1))
            elif var.startswith("`exp"):
                vars_n_DBvars_fetched['exp'].append(var)
                # Extract the variable inside exp and add to 'x'
                inner_var = re.search(r"`exp\(([^)]+)\)", var)
                if inner_var:
                    vars_n_DBvars_fetched['x'].append(inner_var.group(1))
            else:
                vars_n_DBvars_fetched['x'].append(var)

        # Pad shorter lists with None to ensure equal length
        max_length = max(len(v) for v in vars_n_DBvars_fetched.values())
        for key in vars_n_DBvars_fetched:
            vars_n_DBvars_fetched[key].extend([None] * (max_length - len(vars_n_DBvars_fetched[key])))

        df = pd.DataFrame(vars_n_DBvars_fetched)

        # Remove duplicates and clean up DataFrame
        for col in df.columns:
            df[col] = df[col].drop_duplicates().reset_index(drop=True)
        df = df.dropna(axis=0, how='all')

        required_columns = ['x', 'lag', 'log', 'exp']
        for col in required_columns:
            if col not in df.columns:
                df[col] = None

        df = df[['x', 'lag', 'log', 'exp']]  # Reorder columns to match required order
        print("Generated DataFrame:", df)

        # Save cleaned formulas to DOCX
        doc = Document()
        for formula in formulas:
            doc.add_paragraph(formula)
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return df, docx_buffer

    with st.sidebar:
        st.header('File uploads')
        choice_formulas = st.selectbox(label='Formulas', options=['mine', 'default'])
        if choice_formulas == 'mine':
            formulas = st.file_uploader('Upload the docx', type='docx')
            st.session_state['formulas_userinput'] = formulas
            st.session_state['formula_parser_needed'] = True
            with st.expander('Formula (docx) Parser'):
                goparse_button = st.button('Go')
                if goparse_button:
                    formulas_to_parse = st.session_state['formulas_userinput']
                    formulas_parsed, cleaned_formulas = parse_formulas(formulas_to_parse)
                    st.session_state['formula_parser_needed'] = False
                    print(st.session_state['formula_parser_needed'])
                    st.session_state['formulas'] = formulas_parsed
                    st.session_state['formulas_userinput'] = None
                    if not isinstance(formulas_parsed, type(None)):
                        if not formulas_parsed.empty:
                            file_path = "formulas_parsed.csv"  # Fixed file name
                            formulas_parsed.to_csv(file_path, index=False)  # Overwrite if file exists

                            with open(file_path, "rb") as file:
                                st.download_button('Download CSV', data=file, file_name='formulas_parsed.csv',
                                                   mime='text/csv')

                            # DOCX download button
                            st.download_button('Download Clean Formulas', data=cleaned_formulas,
                                               file_name='cleaned_formulas.docx',
                                               mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                            st.write(formulas_parsed)
        elif choice_formulas == 'default':
            formulas = pd.read_excel('streamlit_assets/formulas_mapped.xlsx', engine='openpyxl')
            if 'Unnamed: 0' in formulas.columns:
                formulas.drop('Unnamed: 0', axis=1, inplace=True)
            st.session_state['formulas'] = formulas
            st.session_state['formula_parser_needed'] = False
        else:
            pass

        choice_excel = st.selectbox(label='Mapped Excel Defaults Key', options=['default', 'mine'])
        if choice_excel == 'mine':
            excel_defaults = st.file_uploader('Upload the excel', type='xlsx')
            st.session_state['excel_defaults'] = excel_defaults
            st.session_state['excel_loaded'] = True
        elif choice_excel == 'default':
            excel_defaults = pd.read_excel('streamlit_assets/excel_defaults_mapped.xlsx', sheet_name=None, engine='openpyxl')
            st.session_state['excel_defaults'] = excel_defaults
            st.session_state['excel_loaded'] = True
        else:
            pass

        if st.session_state.get('excel_loaded'):
            country_dictionary = pd.read_csv('streamlit_assets/country_dictionary.csv')
            country_dictionary = country_dictionary._append({'name': 'Türkiye', 'code': 'TR'}, ignore_index=True)
            country_dictionary.drop_duplicates(inplace=True)
            available_codes = excel_defaults.keys()
            country_dictionary = country_dictionary[country_dictionary['code'].isin(available_codes)]
            country_selector = st.selectbox('Country:', options=country_dictionary['name'])

            if country_selector:
                selected_code = country_dictionary[country_dictionary['name'] == country_selector]['code']

                if len(selected_code) == 1:
                    proper_code = str(selected_code.iloc[0])
                    try:
                        excel_defaults_countryfiltered = pd.read_excel('streamlit_assets/excel_defaults_mapped.xlsx',
                                                                       sheet_name=proper_code, engine='openpyxl')
                        st.session_state['excel_defaults_onecountry'] = excel_defaults_countryfiltered
                    except ValueError as e:
                        st.text(e)

        choice_dbkeys = st.selectbox(label='DB Keys', options=['default', 'mine'])
        if choice_dbkeys == 'mine':
            dbkeys = st.file_uploader('Upload db keys (proper_M)', type='csv')
            st.session_state['dbkeys'] = dbkeys
        elif choice_dbkeys == 'default':
            dbkeys = pd.read_csv('streamlit_assets/proper_M.csv')
            st.session_state['dbkeys'] = dbkeys
        else:
            pass

    if all(value is not None for value in
           [st.session_state.get('formulas'),
            st.session_state.get('excel_defaults_onecountry'),
            st.session_state.get('dbkeys')]):

        # if st.sidebar.button('Generate/update'):
        #     if st.session_state['formula_parser_needed'] is True:
        #         st.error('Parse your formulas first.')
        #     else:
        #         st.session_state['sidebar_done'] = True

        if st.sidebar.button('Generate/update'):
            st.session_state['sidebar_done'] = True


if __name__ == '__main__':
    sidepage()
    mainpage()
